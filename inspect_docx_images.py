import docx
from docx import Document

doc_path = "UPDATED CHANGES.docx"
doc = Document(doc_path)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total inline shapes: {len(doc.inline_shapes)}")

# Let's search for actual images by inspecting document body elements.
# In python-docx, inline_shapes is a collection. Let's see if we can find where they are in the paragraphs.
shape_idx = 0
for p_idx, paragraph in enumerate(doc.paragraphs):
    # Find if there are drawings in paragraph run elements
    # and match with actual inline shapes if possible.
    # An inline shape is associated with a drawing element in the run.
    has_real_image = False
    for run in paragraph.runs:
        # Check if the run XML contains a blip (embed image)
        if "w:drawing" in run._r.xml and "r:embed" in run._r.xml:
            has_real_image = True
            
    if has_real_image:
        # Print surrounding context (last 3 paragraphs if available, and current paragraph text)
        print(f"\n--- IMAGE {shape_idx} ---")
        print(f"Located at Paragraph {p_idx}")
        # Print context paragraphs
        start_ctx = max(0, p_idx - 3)
        for i in range(start_ctx, p_idx + 1):
            txt = doc.paragraphs[i].text.strip()
            style = doc.paragraphs[i].style.name
            print(f"  P {i} ({style}): {txt}")
        shape_idx += 1
